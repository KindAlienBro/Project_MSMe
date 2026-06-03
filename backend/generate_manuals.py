import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run 'pip install python-docx' first.")
    exit(1)

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    if bold:
        p.add_run(text).bold = True
    else:
        p.add_run(text)
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def create_admin_manual():
    doc = Document()
    add_heading(doc, 'Project MSMe - Administrator User Manual', 0)
    
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc, 'Welcome to the Administrator User Manual. This guide explains how to manage faculty accounts, approve leaves, and generate AI-powered timetables.')
    
    add_heading(doc, '2. Accessing the Admin Dashboard', 1)
    add_bullet(doc, 'Log in using your Admin credentials.')
    add_bullet(doc, 'Navigate using the left sidebar to access User Management, Timetable Generation, and Leave Requests.')
    
    add_heading(doc, '3. Account Management', 1)
    add_bullet(doc, 'Pending Approvals: View and approve/reject newly registered teacher accounts.')
    add_bullet(doc, 'Active Users: Manage existing accounts. You can promote Teachers to Super Teachers (HODs) or deactivate accounts.')
    
    add_heading(doc, '4. AI Timetable Generation', 1)
    add_bullet(doc, 'Navigate to "Generate Timetable".')
    add_bullet(doc, 'Enter constraints in natural language (e.g., "Mr. Smith needs Friday afternoon off").')
    add_bullet(doc, 'Click "Preview" to let the AI process the timetable (may take up to 2 minutes).')
    add_bullet(doc, 'Click "Apply to Live" to publish the schedule.')
    
    add_heading(doc, '5. Leave & Substitution Management', 1)
    add_bullet(doc, 'View pending leave requests from teachers.')
    add_bullet(doc, 'When you approve a leave, the system automatically assigns substitutes using AI based on free periods and sends them requests.')
    
    doc.save('Admin_User_Manual.docx')
    print("Created Admin_User_Manual.docx")

def create_teacher_manual():
    doc = Document()
    add_heading(doc, 'Project MSMe - Teacher User Manual', 0)
    
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc, 'Welcome to the Teacher User Manual. This guide explains how to view your schedule, take attendance, request leaves, and handle substitutions across both the Web platform and the Flutter Mobile App.')
    
    add_heading(doc, '2. Viewing Your Schedule', 1)
    add_bullet(doc, 'Web: Your dashboard displays today’s classes and your weekly timetable.')
    add_bullet(doc, 'Mobile App: The home screen highlights your next upcoming class.')
    
    add_heading(doc, '3. Taking Attendance', 1)
    add_bullet(doc, 'Click on any Ongoing or Completed class in your schedule.')
    add_bullet(doc, 'A list of students will appear (all marked Present by default).')
    add_bullet(doc, 'Uncheck absent students and click "Submit".')
    add_bullet(doc, 'The system prevents duplicate submissions automatically.')
    
    add_heading(doc, '4. Requesting Leave', 1)
    add_bullet(doc, 'Navigate to the "Leave" tab.')
    add_bullet(doc, 'Select the dates and provide a reason.')
    add_bullet(doc, 'You will be notified once the Admin approves or rejects the request.')
    
    add_heading(doc, '5. Managing Substitutions', 1)
    add_bullet(doc, 'If a colleague is absent, you may receive a substitution request.')
    add_bullet(doc, 'Go to the "Substitutions" tab to view pending requests.')
    add_bullet(doc, 'Click "Accept" or "Decline" before the request expires.')
    
    doc.save('Teacher_User_Manual.docx')
    print("Created Teacher_User_Manual.docx")

def create_student_manual():
    doc = Document()
    add_heading(doc, 'Project MSMe - Student User Manual', 0)
    
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc, 'Welcome to the Student User Manual. This guide explains how to check your timetable and stay updated with real-time class changes.')
    
    add_heading(doc, '2. Student Dashboard Overview', 1)
    add_bullet(doc, 'Live Clock: View the current time and ongoing classes.')
    add_bullet(doc, 'Stats: Track the number of classes, completed sessions, and free periods for the day.')
    
    add_heading(doc, '3. Viewing the Timetable', 1)
    add_bullet(doc, 'Today View: See exactly what classes you have today, the faculty teaching them, and the room numbers.')
    add_bullet(doc, 'Full Week View: Switch tabs to see your entire weekly schedule.')
    
    add_heading(doc, '4. Real-Time Changes & Alerts', 1)
    add_bullet(doc, 'If a teacher is absent, the system will highlight the substituted class in orange.')
    add_bullet(doc, 'You will see a banner alert (e.g., "Mr. Smith is absent. Mrs. Davis will take Physics at 10:00 AM.").')
    add_bullet(doc, 'Cancelled classes appear as Free Periods.')
    
    add_heading(doc, '5. Notifications', 1)
    add_bullet(doc, 'Click the bell icon to view a history of all schedule updates and alerts.')
    
    doc.save('Student_User_Manual.docx')
    print("Created Student_User_Manual.docx")

if __name__ == '__main__':
    create_admin_manual()
    create_teacher_manual()
    create_student_manual()
