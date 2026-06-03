import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed.")
    exit(1)

def add_title(doc, text):
    title = doc.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_step(doc, number, text, bold_intro=""):
    p = doc.add_paragraph()
    if bold_intro:
        p.add_run(f"Step {number}: ").bold = True
        p.add_run(f"{bold_intro} - ").bold = True
        p.add_run(text)
    else:
        p.add_run(f"Step {number}: ").bold = True
        p.add_run(text)
    
def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ 📸 PLACEHOLDER: Please insert screenshot of {text} here ]")
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True
    p.style = 'Intense Quote'

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    if bold:
        p.add_run(text).bold = True
    else:
        p.add_run(text)
    return p

def create_admin_manual():
    doc = Document()
    add_title(doc, 'Project MSMe - Comprehensive Administrator Guide')
    
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc, 'Welcome to the Administrator User Manual. As an administrator, you hold the highest level of access in the Project MSMe ecosystem. This guide provides a detailed, step-by-step walkthrough of all administrative capabilities, including user management, AI-driven timetable generation, and handling faculty leaves.')
    
    add_heading(doc, '2. Accessing the Admin Dashboard', 1)
    add_paragraph(doc, 'The Admin Dashboard is your central hub for all operations. From here, you can view the health of the system and access all primary modules.')
    add_step(doc, 1, 'Navigate to the web portal login page.', 'Open your browser')
    add_step(doc, 2, 'Enter your administrative email and password.', 'Log In')
    add_step(doc, 3, 'Upon successful authentication, you will be redirected to the main Admin Dashboard.', 'Dashboard View')
    add_placeholder(doc, 'Main Admin Dashboard showing the left sidebar and summary widgets')
    
    add_heading(doc, '3. Account Management', 1)
    add_paragraph(doc, 'For security purposes, when a new teacher registers on the platform, their account must be manually approved by an administrator before they can access the system.')
    
    add_heading(doc, '3.1 Approving New Teachers', 2)
    add_step(doc, 1, 'Click on the "Users" or "Account Management" tab in the left sidebar.', 'Navigate to Users')
    add_step(doc, 2, 'Look for the section titled "Pending Approvals". This list contains all newly registered accounts.', 'Locate Pending Accounts')
    add_step(doc, 3, 'Review the teacher\'s name, department, and email address to verify their identity.', 'Review Details')
    add_step(doc, 4, 'Click the green "Approve" button next to their name. The teacher will immediately gain access to the platform.', 'Approve Account')
    add_placeholder(doc, 'Pending Approvals list with the Approve/Reject buttons visible')
    
    add_heading(doc, '3.2 Managing Active Users', 2)
    add_step(doc, 1, 'Scroll down to the "Active Users" table.', 'View Active Users')
    add_step(doc, 2, 'To promote a teacher to a Head of Department (HOD), click the "Promote to Super Teacher" button. This grants them department-level admin rights.', 'Promote Faculty')
    add_step(doc, 3, 'To remove access from a teacher who has left the institution, click the red "Deactivate" button. This preserves their historical data but prevents login.', 'Deactivate Account')
    add_placeholder(doc, 'Active Users table showing the Promote and Deactivate buttons')

    add_heading(doc, '4. AI Timetable Generation', 1)
    add_paragraph(doc, 'Project MSMe uses a powerful AI model to automatically generate conflict-free schedules for the entire institution based on your natural language constraints.')
    
    add_heading(doc, '4.1 Creating a New Timetable', 2)
    add_step(doc, 1, 'Click on "Generate Timetable" in the sidebar.', 'Open Generator')
    add_step(doc, 2, 'In the text box provided, type any specific constraints you have in plain English. For example: "Ensure Mr. Smith has no classes on Friday afternoons because he is on half-day leave." or "Do not schedule consecutive Math classes for Section 3A."', 'Enter Constraints')
    add_placeholder(doc, 'Timetable Generator screen showing the natural language text box')
    add_step(doc, 3, 'Click the "Preview" button. The AI solver will begin processing. Note: This process requires heavy computation and may take up to 2 minutes.', 'Run AI Solver')
    add_step(doc, 4, 'Once complete, the system will display a visual grid of the newly generated schedule. Review it to ensure your constraints were met.', 'Review Preview')
    add_placeholder(doc, 'The generated timetable preview grid')
    add_step(doc, 5, 'If you are satisfied with the schedule, click "Apply to Live". This will instantly update the schedules for all teachers and students on their dashboards and mobile apps.', 'Publish Timetable')

    add_heading(doc, '5. Leave & Substitution Management', 1)
    add_paragraph(doc, 'When a teacher applies for leave, the system intelligently calculates the optimal substitute teachers for their classes and handles the assignment automatically.')
    add_step(doc, 1, 'Go to the "Leave Requests" tab.', 'View Leaves')
    add_step(doc, 2, 'Review the pending leave requests, including the dates and reason provided by the teacher.', 'Review Request')
    add_step(doc, 3, 'Click "Approve". Behind the scenes, the AI will find teachers who are free during those periods and send them a substitution request.', 'Approve & Trigger AI')
    add_placeholder(doc, 'Leave request table showing the Approve button')
    
    doc.save('Admin_User_Manual_Detailed.docx')
    print("Created Admin_User_Manual_Detailed.docx")


def create_teacher_manual():
    doc = Document()
    add_title(doc, 'Project MSMe - Comprehensive Teacher Guide')
    
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc, 'Welcome to the Teacher User Manual. This guide provides detailed, step-by-step instructions on how to manage your daily academic tasks, including viewing your schedule, taking attendance, requesting leaves, and handling substitutions.')
    
    add_heading(doc, '2. Viewing Your Schedule', 1)
    add_paragraph(doc, 'Your schedule is available on both the Web Portal and the Flutter Mobile App.')
    
    add_heading(doc, '2.1 Web Portal Dashboard', 2)
    add_step(doc, 1, 'Log in to the web portal.', 'Log In')
    add_step(doc, 2, 'On the main dashboard, you will see a "Today\'s Schedule" card. This shows a timeline of all your classes for the current day.', 'Today\'s View')
    add_placeholder(doc, 'Teacher Web Dashboard showing Today\'s Schedule timeline')
    add_step(doc, 3, 'To view your entire week, click the "Full Week" toggle button at the top of the dashboard. This displays a comprehensive grid of all your classes.', 'Weekly View')
    add_placeholder(doc, 'Teacher Web Dashboard showing the Full Week grid')
    
    add_heading(doc, '2.2 Mobile App Schedule', 2)
    add_step(doc, 1, 'Open the Project MSMe app on your mobile device.', 'Launch App')
    add_step(doc, 2, 'The home screen prominently displays your next upcoming class with a countdown timer.', 'Next Class View')
    add_placeholder(doc, 'Mobile App home screen showing the upcoming class card')

    add_heading(doc, '3. Taking Attendance', 1)
    add_paragraph(doc, 'Attendance can be taken seamlessly from the web or your mobile phone. The system is designed to prevent duplicate entries.')
    
    add_step(doc, 1, 'On your dashboard (web or mobile), click on the specific class session you are currently teaching.', 'Select Class')
    add_step(doc, 2, 'Click the "Take Attendance" button.', 'Open Attendance Portal')
    add_placeholder(doc, 'The "Take Attendance" button on the class details card')
    add_step(doc, 3, 'A list of all enrolled students will appear. By default, every student is marked as "Present".', 'View Roster')
    add_step(doc, 4, 'Scroll through the list and UNCHECK the box (or toggle off) next to the names of students who are absent.', 'Mark Absentees')
    add_placeholder(doc, 'Attendance list with some students unchecked')
    add_step(doc, 5, 'Scroll to the bottom and click the blue "Submit Attendance" button.', 'Submit')
    add_paragraph(doc, 'Note: If you or a co-teacher have already submitted attendance for this exact session, the system will show an orange "Already Submitted" warning and will not overwrite the existing data.')

    add_heading(doc, '4. Requesting Leave', 1)
    add_step(doc, 1, 'Click on the "Leaves" tab in your navigation menu.', 'Open Leaves')
    add_step(doc, 2, 'Click the "New Leave Request" button.', 'Create Request')
    add_step(doc, 3, 'Select your start date and end date from the calendar widget.', 'Select Dates')
    add_step(doc, 4, 'Type a brief reason for your absence (e.g., "Medical emergency" or "Attending seminar").', 'Provide Reason')
    add_step(doc, 5, 'Click "Submit Request". You will receive a notification when the HOD or Admin approves it.', 'Submit')
    add_placeholder(doc, 'The Leave Request form modal')

    add_heading(doc, '5. Managing Substitutions', 1)
    add_paragraph(doc, 'If a colleague is absent and you have a free period, the AI may assign you as a substitute. You have the right to accept or decline this request.')
    add_step(doc, 1, 'When a substitution is requested, a red badge will appear on your notification bell icon.', 'Check Notifications')
    add_step(doc, 2, 'Navigate to the "Substitutions" tab from the sidebar.', 'Open Substitutions')
    add_step(doc, 3, 'You will see a pending request showing the subject, time, section, and the absent teacher\'s name.', 'Review Request')
    add_placeholder(doc, 'Pending substitution request card')
    add_step(doc, 4, 'Click the green "Accept" button to take the class, or the red "Decline" button if you are unavailable.', 'Respond')
    add_paragraph(doc, 'Note: Requests expire automatically once the class time passes. Expired requests will be greyed out and cannot be interacted with.')
    
    doc.save('Teacher_User_Manual_Detailed.docx')
    print("Created Teacher_User_Manual_Detailed.docx")

def create_student_manual():
    doc = Document()
    add_title(doc, 'Project MSMe - Comprehensive Student Guide')
    
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc, 'Welcome to the Student User Manual. The student portal is designed to keep you perfectly informed about your daily schedule, real-time timetable changes, and attendance records.')
    
    add_heading(doc, '2. The Student Dashboard', 1)
    add_step(doc, 1, 'Log in to your student account.', 'Log In')
    add_step(doc, 2, 'At the top of the dashboard, you will see a live clock and date widget that adjusts to your local time.', 'Live Clock')
    add_placeholder(doc, 'Student Dashboard header with the Live Clock and Stats widgets')
    add_step(doc, 3, 'Below the clock, check the statistics widgets to see exactly how many classes you have today, how many are completed, and if you have any free periods.', 'Daily Stats')
    
    add_heading(doc, '3. Viewing Your Timetable', 1)
    add_step(doc, 1, 'By default, the dashboard shows "Today\'s Schedule", listing your classes in chronological order.', 'Today\'s Classes')
    add_step(doc, 2, 'Look at each class card to see the subject, the faculty teaching it, the room number, and the exact time.', 'Class Details')
    add_placeholder(doc, 'List of today\'s classes on the student dashboard')
    add_step(doc, 3, 'To see your schedule for tomorrow or the rest of the week, click the "Full Week" toggle button. This will display a comprehensive grid of all your classes.', 'Weekly Grid')
    
    add_heading(doc, '4. Tracking Real-Time Changes', 1)
    add_paragraph(doc, 'Project MSMe updates your schedule instantly if a teacher is absent.')
    add_step(doc, 1, 'Look for banner alerts at the top of your dashboard. If a substitution occurs, you will see a message like "Mr. Smith is absent. Mrs. Davis will take Physics at 10:00 AM."', 'Check Alerts')
    add_placeholder(doc, 'Orange alert banner indicating a schedule change')
    add_step(doc, 2, 'In your class list, substituted classes will have an orange "Substituted" tag next to the subject name. The name of the new faculty member will be displayed.', 'Identify Substitutions')
    add_step(doc, 3, 'If a class is completely cancelled and no substitute is found, it will be marked explicitly as a "Free Period" in your timeline.', 'Free Periods')

    add_heading(doc, '5. Viewing Notifications', 1)
    add_step(doc, 1, 'Click the bell icon (🔔) in the top right corner of your screen.', 'Open Notifications')
    add_step(doc, 2, 'A dropdown will appear showing a history of all announcements, schedule changes, and alerts.', 'Review History')
    add_placeholder(doc, 'The notifications dropdown menu')

    doc.save('Student_User_Manual_Detailed.docx')
    print("Created Student_User_Manual_Detailed.docx")

if __name__ == '__main__':
    create_admin_manual()
    create_teacher_manual()
    create_student_manual()
